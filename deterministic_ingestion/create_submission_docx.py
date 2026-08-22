#!/usr/bin/env python3
"""Generate the Assignment 2 Word submission from current workspace evidence."""
import json
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).parent.parent
RAW = ROOT / 'second_brain' / 'raw'
NORMALIZED = ROOT / 'second_brain' / 'normalized'
ASSETS = ROOT / 'submission_evidence'
OUT = ROOT / 'northstar_assignment2_submission.docx'


def files(directory):
    return sorted(directory.glob('*')) if directory.exists() else []


def read_json(path):
    try:
        return json.loads(path.read_text(encoding='utf-8'))
    except Exception:
        return {}


def set_cell_shading(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    shade = OxmlElement('w:shd')
    shade.set(qn('w:fill'), fill)
    props.append(shade)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ''
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_table(document, rows, widths=None):
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for index, value in enumerate(rows[0]):
        set_cell_text(table.rows[0].cells[index], value, bold=True, color=(255, 255, 255))
        set_cell_shading(table.rows[0].cells[index], '17324D')
    for row in rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    document.add_paragraph()
    return table


def heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    paragraph.runs[0].font.color.rgb = RGBColor(23, 50, 77)
    return paragraph


def code(document, text):
    paragraph = document.add_paragraph()
    paragraph.style = 'No Spacing'
    paragraph.paragraph_format.left_indent = Inches(0.2)
    run = paragraph.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    return paragraph


def main():
    raw_jobs = files(RAW / 'jobs')
    raw_github = files(RAW / 'github')
    raw_learning = files(RAW / 'learning')
    norm_jobs = files(NORMALIZED / 'jobs')
    norm_github = files(NORMALIZED / 'github')
    norm_learning = files(NORMALIZED / 'learning')
    scores = files(NORMALIZED / 'jobs_scores')
    zapier_raw = next((p for p in raw_jobs if 'senior-data-engineer' in p.name), raw_jobs[0] if raw_jobs else None)
    zapier_norm = next((p for p in norm_jobs if 'senior-data-engineer' in p.name), norm_jobs[0] if norm_jobs else None)
    score_sample = next((p for p in scores if 'github_repo' in p.name), scores[0] if scores else None)
    job_data = read_json(zapier_norm) if zapier_norm else {}
    score_data = read_json(score_sample) if score_sample else {}
    raw_text = zapier_raw.read_text(encoding='utf-8', errors='ignore') if zapier_raw else ''

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    normal = doc.styles['Normal']
    normal.font.name = 'Aptos'
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)

    title = doc.add_heading('North Star\nSecond Brain', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(23, 50, 77)
    subtitle = doc.add_paragraph('Assignment 2 submission | Evidence captured 22 August 2026')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('A deterministic personal information pipeline that collects real job, learning, RSS, and GitHub sources, then ranks them against one personally chosen career goal.').style = 'Intense Quote'
    doc.add_paragraph('Applicant goal: move from Data Engineer to Generative AI Engineer by April 2027, using evidence from real work, learning, and opportunities.')

    heading(doc, '1. North Star statement')
    doc.add_paragraph('By April 2027, starting from my current role as Data Engineer, I will be working as a Generative AI Engineer at a company, verified by an employment start date and an official offer letter, so I can make production AI engineering my next career focus.').style = 'Intense Quote'
    heading(doc, 'Four parts a stranger can check', 2)
    add_table(doc, [
        ['Part', 'What it says'],
        ['Current state', 'I am currently a Data Engineer.'],
        ['Time-bound goal', 'The deadline is April 2027.'],
        ['Provable activity', 'Employment as a Generative AI Engineer, checked using an employment start date and official offer letter.'],
        ['Reason', 'I want production AI engineering to become my next career focus.'],
    ], [1.3, 5.9])

    heading(doc, '2. Source list')
    doc.add_paragraph('Sources connected to this project. The two strongest signals are marked [MOST SIGNAL].')
    add_table(doc, [
        ['Source', 'What enters the brain', 'Signal judgment'],
        ['LinkedIn Job Alerts via IMAP [MOST SIGNAL]', 'Real job-alert emails, including multiple jobs per digest.', 'Primary opportunity signal.'],
        ['Zapier structured job email [MOST SIGNAL]', 'Structured title, company, location, URL, date, tags, and responsibilities.', 'Primary clean job signal.'],
        ['GitHub repository search', 'GenAI repositories sorted by stars or recent updates.', 'Practical skill and trend signal.'],
        ['RSS feed', 'Public GenAI engineering articles and posts.', 'Knowledge signal; not employment proof.'],
        ['LMS folder', 'Course notes and exported learning material.', 'Personal learning signal.'],
        ['Consciously cut', 'Personalized LinkedIn home feed: no stable official RSS route and mixed signal. Google security and Pipedream notices: operational mail, not career inputs.', 'Deliberately excluded; job alerts and public GenAI feeds remain.'],
    ], [2.0, 3.5, 1.7])

    heading(doc, '3. Second brain folder and real items')
    doc.add_paragraph('The project stores raw Markdown separately from normalized and synthesized outputs. Current evidence counts:')
    add_table(doc, [
        ['Folder', 'Items', 'Purpose'],
        ['second_brain/raw/jobs/', len(raw_jobs), 'Raw email and RSS job material.'],
        ['second_brain/raw/github/', len(raw_github), 'Raw GitHub repository search results.'],
        ['second_brain/raw/learning/', len(raw_learning), 'Raw LMS Markdown notes.'],
        ['second_brain/normalized/jobs/', len(norm_jobs), 'Structured job records.'],
        ['second_brain/normalized/github/', len(norm_github), 'Structured repository records.'],
        ['second_brain/normalized/learning/', len(norm_learning), 'Structured learning records.'],
        ['second_brain/normalized/jobs_scores/', len(scores), 'LLM ranking output.'],
    ], [3.0, 0.7, 3.5])
    doc.add_paragraph('Folder tree:')
    code(doc, 'second_brain/\n  raw/\n    jobs/\n    github/\n    learning/\n  normalized/\n    jobs/\n    github/\n    learning/\n    jobs_scores/')
    doc.add_paragraph('Evidence capture: generated from the current repository folders.')
    doc.add_picture(str(ASSETS / 'folder_structure.png'), width=Inches(6.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Examples: LinkedIn Generative AI Engineer alerts, Senior Data Engineer from Zapier, Scaling LLM-based Ranking Systems with SGLang, Microsoft Generative AI for Beginners, LangChain, RAGFlow, vLLM, and three LMS learning notes.')

    heading(doc, '4. Personalization engine')
    doc.add_paragraph('The engine uses one scoring prompt in deterministic_ingestion/SCORE_PROMPT.md. It asks GROQ to score every normalized item against the North Star using five weighted dimensions.')
    add_table(doc, [
        ['Dimension', 'Weight', 'Meaning'],
        ['Alignment', '40', 'Direct connection to becoming a Generative AI Engineer.'],
        ['Impact on goal', '20', 'Employment or strong portfolio value.'],
        ['Feasibility / proximity', '15', 'Achievability from the current Data Engineer role.'],
        ['Skill match', '15', 'LLM, RAG, inference, MLOps, and related skills.'],
        ['Urgency / timing', '10', 'Timeliness and opportunity window.'],
    ], [2.0, 0.7, 4.5])
    doc.add_paragraph('The LLM returns JSON with score, breakdown, confidence, and notes. Jobs are judged for role fit, GitHub repositories for practical implementation, LMS items for hands-on learning, and RSS for learning and trend signals.')
    heading(doc, 'Before and after example', 2)
    doc.add_paragraph('Before: raw incoming Zapier email')
    code(doc, ' '.join(raw_text.split())[:1200])
    doc.add_paragraph('After: normalized record')
    after = {k: job_data.get(k) for k in ('source', 'email_type', 'title', 'company', 'location', 'url', 'date_posted', 'tags', 'skills', 'responsibilities')}
    code(doc, json.dumps(after, indent=2)[:1800])
    doc.add_paragraph('Filtered/scored output example')
    code(doc, json.dumps(score_data, indent=2)[:1000])
    doc.add_paragraph('Code proof: the scorer scans all three normalized source folders:')
    code(doc, "NORM_DIRS = [NORM_ROOT / name for name in ('jobs', 'github', 'learning')]\nfiles = sorted(path for norm_dir in NORM_DIRS\n                for path in glob(str(norm_dir / '*.json')))\noutput = call_llm(build_prompt(north_star, item))\nparse_and_save(output, item_id)")
    doc.add_picture(str(ASSETS / 'pipeline_before_after.png'), width=Inches(6.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading(doc, '5. Ingestion mechanism')
    doc.add_paragraph('The pipeline runs manually or through GitHub Actions. The weekly workflow runs every Monday at 06:00 UTC and also supports workflow_dispatch.')
    add_table(doc, [
        ['Step', 'Implementation evidence'],
        ['GitHub repositories', 'ingest_github.py searches GenAI repositories and writes raw Markdown.'],
        ['RSS knowledge', 'ingest_linkedin_rss.py reads configured RSS/Atom feeds and deduplicates entries.'],
        ['Email jobs', 'ingest_email_imap.py fetches unseen messages and filters unrelated operational mail.'],
        ['LMS', 'ingest_lms_folder.py imports files from inbox_lms/.'],
        ['Normalization', 'normalize_job.py handles jobs; normalize_sources.py handles GitHub and LMS.'],
        ['Scoring', 'score_with_llm.py scans all normalized folders and calls GROQ.'],
    ], [1.6, 5.6])
    doc.add_paragraph('Workflow evidence: .github/workflows/weekly-ingest.yml. Manual schedule evidence: deterministic_ingestion/cron_example.txt.').style = 'Intense Quote'
    code(doc, '0 * * * * ingest_linkedin_rss.py\n15 * * * * ingest_github.py\n30 * * * * ingest_email.py')
    doc.add_paragraph('Code proof: workflow command sequence:')
    code(doc, 'python deterministic_ingestion/ingest_email_imap.py\npython deterministic_ingestion/ingest_linkedin_rss.py\npython deterministic_ingestion/ingest_github.py\npython deterministic_ingestion/ingest_lms_folder.py\npython deterministic_ingestion/normalize_job.py\npython deterministic_ingestion/normalize_sources.py\npython deterministic_ingestion/score_with_llm.py')
    doc.add_picture(str(ASSETS / 'workflow_evidence.png'), width=Inches(6.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('The raw records remain Markdown. The normalizers create structured JSON. Only the final scoring step calls GROQ, and its output is stored separately in jobs_scores.')

    doc.save(OUT)
    print(OUT)


if __name__ == '__main__':
    main()
